from deep_translator import GoogleTranslator
import threading
import docx
from tqdm import tqdm


def checkInteger(string):
    try:
        int(string)
        return True
    except ValueError:
        return False


def translateParagraph(doc, start, end, bar):
    for i in range(start, end):
        paragraph = doc.paragraphs[i].text
        if(paragraph != "" and checkInteger(paragraph) == False):
            translated = GoogleTranslator(source='en', target='pt').translate(doc.paragraphs[i].text)
            doc.paragraphs[i].text = translated
            bar.update(1)

def translateDocx(threadsToUse):
    doc = docx.Document("output.docx")

    print(len(doc.paragraphs))
    threads = []

    paragraphs = len(doc.paragraphs)
    dividedParagraphs = paragraphs//threadsToUse

    paragraphsExcluded = 0
    for k, p in enumerate(doc.paragraphs):
        if p.text != "" and not checkInteger(p.text):
            paragraphsExcluded = paragraphsExcluded + 1

    progressBar = tqdm(leave=False, total=paragraphsExcluded, desc=f'[Paragraphs translated]', unit="pg")

    for i in range(threadsToUse):
        if i == 0:
            start = 0
            end = dividedParagraphs
        else:
            start = dividedParagraphs*i
            end = start + dividedParagraphs + 1

        if i == threadsToUse-1:
            end = paragraphs

        thread = threading.Thread(target=translateParagraph, args=(doc, start, end, progressBar))
        threads.append(thread)
        thread.start()

    for key, thread in enumerate(threads):
        thread.join()

    progressBar.close()


    doc.save("outputTranslated.docx")
